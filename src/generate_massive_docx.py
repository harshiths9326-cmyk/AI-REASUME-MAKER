import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_borders(doc):
    for sec in doc.sections:
        sectPr = sec._sectPr
        # Remove existing pgBorders if any
        existing_borders = sectPr.xpath('.//w:pgBorders')
        if existing_borders:
            for eb in existing_borders:
                sectPr.remove(eb)
                
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '24') # 3pt
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')
            pgBorders.append(border)
            
        sectPr.append(pgBorders)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 0:
            run.font.size = Pt(24)
        elif level == 1:
            run.font.size = Pt(18)
        else:
            run.font.size = Pt(14)
    return h

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    return p

content_sections = [
    {
        "title": "1. ABSTRACT & INTRODUCTION",
        "content": [
            "The AI Resume Maker is an intelligent, web-based application designed to automate and professionalize the process of resume creation. Developing a high-quality resume is often a hurdle for students and freshers due to formatting complexities, lack of understanding of industry standards, and the sheer time required to align elements in traditional word processors. This project comprehensively addresses these multifaceted challenges by providing a highly intuitive, user-friendly platform built upon modern web technologies including Next.js, TypeScript, and Supabase. The system features dynamic template switching, real-time preview rendering, secure cloud storage for data persistence, and high-fidelity PDF generation capabilities. By leveraging a modular component-based architecture and cloud-native services, the application ensures horizontal scalability, optimal performance metrics, and accessibility across all standard devices and browser environments. This report thoroughly documents the end-to-end software development life cycle, from initial requirement analysis and abstract system design to final implementation, testing, and deployment.",
            "In the modern hyper-competitive job market, a resume functions as significantly more than just a chronological list of academic and professional qualifications; it is a critical marketing tool that represents a candidate's professional identity to prospective employers and Applicant Tracking Systems (ATS). However, the manual creation and formatting of resumes remain a notoriously tedious and error-prone task for many individuals. Users frequently struggle with traditional desktop word processors, battling alignment issues, inconsistent margin spacing, and the challenge of selecting appropriate, modern layouts that stand out without violating professional norms. The AI Resume Maker fundamentally disrupts this outdated paradigm by offering a structured, data-driven, and automated approach to resume building.",
            "Traditional methods of resume creation involve using desktop software which, while undeniably powerful for general document editing, requires significant manual effort and domain-specific formatting knowledge. Beginners often face inconsistent spacing, difficulty in logically organizing technical projects, and a distinct lack of professional-grade templates out of the box. Furthermore, there is no centralized, cloud-accessible place to store, manage, and update multiple tailored versions of a resume for different job applications. Existing online builders often hide their best professional templates behind exorbitant paywalls, lack the responsive design needed for seamless mobile use, or generate PDFs that break ATS parsers. The proposed AI Resume Maker provides a full-stack, universally accessible solution that completely decouples the user's raw content from the complex visual presentation layer."
        ],
        "image": r"C:\Users\Bhoomika\.gemini\antigravity\brain\7729b00b-8521-4b12-ad09-07549338132e\landing_page_screenshot_1778698938455.png"
    },
    {
        "title": "2. SYSTEM ARCHITECTURE AND TECHNOLOGY STACK",
        "content": [
            "The architectural foundation of the AI Resume Maker is constructed upon a robust, modern 3-tier architecture, ensuring a strict separation of concerns, enhanced maintainability, and seamless scalability. The client layer is powered by React and the Next.js framework, leveraging both Server-Side Rendering (SSR) for initial load performance and Client-Side Routing for a fluid, SPA-like user experience. The business logic is encapsulated within Next.js API Routes, which act as serverless functions, securely processing client requests without the overhead of maintaining a traditional monolithic server. The data persistence layer is handled by Supabase, a highly scalable, open-source alternative to Firebase that provides a managed PostgreSQL database, robust authentication mechanisms, and real-time data subscription capabilities.",
            "Next.js was explicitly chosen over traditional React SPAs due to its superior handling of route-based code splitting and built-in optimization for web vitals. By utilizing Next.js's App Router, the application benefits from nested layouts and React Server Components, significantly reducing the JavaScript bundle size shipped to the client browser. This is particularly crucial for a resume builder, where the user expects near-instantaneous feedback as they input their data. Furthermore, Tailwind CSS was integrated as the primary styling solution. By employing a utility-first CSS framework, the development team was able to rapidly prototype UI components without the cognitive load of managing massive, cascading CSS files, ensuring that the application's aesthetic remains consistent across all 12+ dynamically rendered templates.",
            "For the backend infrastructure, Supabase provides an unparalleled developer experience by exposing a RESTful API directly over the underlying PostgreSQL database. This eliminates the need to write boilerplate CRUD (Create, Read, Update, Delete) endpoints. Security is enforced at the database level using PostgreSQL's Row Level Security (RLS) policies. These policies cryptographically guarantee that an authenticated user can only query, mutate, or delete resume records that explicitly belong to their unique User ID, thus preventing unauthorized data access or cross-tenant data leakage. The use of the JSONB data type in PostgreSQL allows the application to store the deeply nested resume structure (containing arrays of experiences, education, and skills) in a schema-less format, providing immense flexibility for future feature additions without requiring complex database migrations."
        ]
    },
    {
        "title": "3. USER EXPERIENCE (UX) AND TEMPLATE DESIGN: MODERN AI",
        "content": [
            "The user experience (UX) paradigm of the AI Resume Maker centers entirely around the concept of 'Information at a Glance' and immediate visual feedback. The primary editor interface adopts a split-pane layout: the left pane houses a vertically scrollable, highly organized form for data entry, while the right pane displays a high-fidelity, real-time preview of the generated resume. This dual-pane approach eliminates the jarring context switching inherent in traditional web forms where users must navigate to a separate 'preview' page to see their changes. As the user types into the input fields, React's state management immediately propagates the changes down the component tree, updating the preview pane in less than 50 milliseconds, creating a deeply engaging and responsive application feel.",
            "The 'Modern AI' template was meticulously engineered specifically for candidates in technical, engineering, and analytical roles. This design philosophy deliberately eschews overly complex graphical elements in favor of stark, minimalist aesthetics that prioritize content readability and structural clarity above all else. The layout utilizes a clean, stark white background contrasted with strong, primary color accents (typically a deep, professional blue or a sharp, slate grey) applied strategically to section headers and horizontal dividers. The header section is geometrically justified, ensuring that the candidate's name and contact information form a visually anchoring block at the top of the document.",
            "Typography in the Modern AI template relies exclusively on modern, sans-serif web fonts such as 'Inter' or 'Roboto'. These fonts are specifically chosen for their excellent legibility on both high-resolution digital displays and in printed formats. Section titles are rendered in bold, uppercase lettering with slight letter-spacing to create a strong visual hierarchy, guiding the recruiter's eye effortlessly down the page. The experience and education sections are formatted using a rigid grid system, ensuring that dates are perfectly right-aligned while job titles and company names remain flush left. This strict adherence to alignment creates a sense of meticulousness and professionalism, subtle cues that are highly valued in technical recruitment processes."
        ],
        "image": r"C:\Users\Bhoomika\.gemini\antigravity\brain\7729b00b-8521-4b12-ad09-07549338132e\template_modern_ai_1778698574877.png"
    },
    {
        "title": "4. TEMPLATE ANALYSIS: EXECUTIVE PROFESSIONAL",
        "content": [
            "In stark contrast to the minimalist approach of the Modern AI design, the 'Executive Professional' template caters explicitly to seasoned professionals, management candidates, and individuals in traditional corporate environments such as finance, law, and business administration. The structural architecture of this template is predicated on a classic, highly conservative two-column layout. This specific structural choice allows for the maximal density of information without overwhelming the reader, which is a critical requirement for candidates possessing decades of extensive work history and numerous accolades.",
            "The aesthetic foundation of the Executive Professional template relies heavily on traditional serif typography, specifically utilizing font families like 'Merriweather' or 'Georgia'. Serif fonts have long been established as the standard in formal business communication, conveying a sense of authority, established credibility, and serious academic rigor. The typographic scale is tightly controlled, with subtle variations in font weight and italics used to differentiate between job titles, company names, and dates, rather than relying on disruptive color changes or graphical icons.",
            "The visual hierarchy in this template is driven by horizontal rules and generous, calculated whitespace between major sections. The primary column is dedicated to the core chronological work experience and comprehensive educational background, while the narrower secondary column efficiently houses contact information, technical proficiencies, and professional certifications. This deliberate separation ensures that a recruiter can instantly locate the candidate's hard skills without having to parse through dense paragraphs of experiential descriptions. The final generated PDF maintains exact millimeter precision, ensuring that the printed document feels like a premium, professionally typeset curriculum vitae."
        ],
        "image": r"C:\Users\Bhoomika\.gemini\antigravity\brain\7729b00b-8521-4b12-ad09-07549338132e\template_corporate_professional_1778698667541.png"
    },
    {
        "title": "5. TEMPLATE ANALYSIS: MODERN CREATIVE",
        "content": [
            "The 'Modern Creative' template was developed to fulfill the specific needs of users in creative industries—such as graphic design, UI/UX architecture, digital marketing, and content creation—where visual differentiation and personal branding are just as critical as the textual content of the resume. The architectural layout breaks away from traditional monochrome constraints by introducing a prominent, colored vertical sidebar that runs the entire length of the document. This bold structural element immediately captures attention and provides a distinct, memorable visual identity.",
            "The sidebar serves a dual functional and aesthetic purpose. Aesthetically, it acts as a canvas for the user to inject their personal brand color, typically defaulting to a sophisticated dark slate (bg-[#2a303c]) or a vibrant, modern hue. Functionally, the sidebar effectively segregates secondary information—such as contact details, social media links (LinkedIn, GitHub, Portfolio), education summaries, and graphical skill representations—away from the primary narrative of the work experience. This creates a highly scannable document where the reader's eye is naturally drawn to the critical achievements located in the expansive, white main content area.",
            "Typography within the Modern Creative template is highly dynamic, often pairing a bold, geometric sans-serif font for headers (like 'Montserrat' or 'Poppins') with a highly readable, softer sans-serif for the body text. This contrast creates a modern, edgy, and forward-thinking aesthetic. Furthermore, this template incorporates subtle UI elements within the document itself, such as pill-shaped badges for technical skills or minimalist progress bars to visually indicate language proficiency. These graphical additions not only break up textual monotony but also demonstrate the candidate's understanding of modern interface design principles directly through their resume presentation."
        ],
        "image": r"C:\Users\Bhoomika\.gemini\antigravity\brain\7729b00b-8521-4b12-ad09-07549338132e\template_modern_creative_1778698741933.png"
    },
    {
        "title": "6. TEMPLATE ANALYSIS: MODERN TECH",
        "content": [
            "The 'Modern Tech' template represents a bold departure from conventional resume design, specifically engineered to resonate with hiring managers in the software development, cybersecurity, and cutting-edge technology sectors. Drawing heavy inspiration from Integrated Development Environments (IDEs) and developer-centric interfaces, this template employs a high-contrast, 'dark mode' aesthetic. The foundational background is a deep, immersive charcoal or near-black, while the text is rendered in stark white and vibrant, syntax-highlighting-inspired accent colors (such as neon green, cyber blue, or alert yellow).",
            "This template is not merely a palette swap; its entire typographic and structural foundation is optimized for the tech industry. It heavily utilizes monospaced fonts, such as 'Fira Code' or 'JetBrains Mono', for section headers and technical skill listings. This deliberate typographic choice instantly signals to the reader that the candidate is deeply entrenched in the coding ecosystem. The layout is highly modular, often presenting technical projects and GitHub contributions with the same level of prominence as formal employment history, reflecting the reality that in modern software engineering, open-source contributions are often as valuable as corporate tenure.",
            "Beyond its striking visual appearance, the Modern Tech template is meticulously engineered to ensure perfect Applicant Tracking System (ATS) compatibility despite its unconventional look. The underlying DOM structure maintains strict semantic HTML ordering, ensuring that parsing algorithms can accurately extract the text without being confused by the visual styling. The PDF generation engine translates this dark theme flawlessly into vector graphics, ensuring that even if a recruiter chooses to print the document, the high contrast remains legible and striking, making the candidate utterly unforgettable in a sea of standard white-paper resumes."
        ],
        "image": r"C:\Users\Bhoomika\.gemini\antigravity\brain\7729b00b-8521-4b12-ad09-07549338132e\template_modern_tech_1778698775717.png"
    },
    {
        "title": "7. TEMPLATE ANALYSIS: ELEGANT SERIF",
        "content": [
            "The 'Elegant Serif' template is the epitome of refined, classic document design, tailored specifically for academia, publishing, executive leadership, and highly formal corporate sectors. It consciously rejects the hyper-modern, colorful trends seen in other templates, opting instead for a timeless, monochromatic elegance that speaks to stability, deep expertise, and traditional professionalism. The entire structural integrity of this design is based on classical typographic principles, emphasizing the rhythm and flow of the text rather than relying on graphical embellishments.",
            "Central to this template's success is its exclusive use of high-quality serif typography, most notably 'Merriweather' or 'Playfair Display'. These fonts are characterized by high contrast between thick and thin strokes and delicate, classical serifs, which guide the reader's eye smoothly across long lines of text. The typographic hierarchy is established through a masterful use of sizing, small caps, and subtle italicization. Section headers are often centered, creating a formal, almost monumental symmetry to the document, while job titles and company names are carefully balanced to maintain visual equilibrium.",
            "Whitespace, or 'negative space', is treated as an active design element in the Elegant Serif template. Margins are notably wider, and line spacing (leading) is intentionally generous, ensuring that the document never feels cramped or overwhelming, even when packed with decades of experience. This breathing room allows the text to stand on its own merits. The resulting PDF is indistinguishable from a document typeset by a professional designer using advanced publishing software, ensuring that the candidate's application is perceived with the utmost respect and gravitas."
        ],
        "image": r"C:\Users\Bhoomika\.gemini\antigravity\brain\7729b00b-8521-4b12-ad09-07549338132e\template_elegant_serif_clean_1778698816640.png"
    },
    {
        "title": "8. PDF GENERATION ENGINE AND BROWSER DOM MANIPULATION",
        "content": [
            "The core technical challenge of the AI Resume Maker lies in its ability to convert complex, dynamically generated HTML and CSS into a static, perfectly formatted, and highly compressed PDF document. This challenge is met through a sophisticated implementation of the `react-to-print` library, combined with meticulous, print-specific CSS media queries. The architecture dictates that the 'Preview' pane is not just a visual facsimile; it is the exact Document Object Model (DOM) node that will be captured and transformed into the final PDF output. This ensures absolute parity between what the user sees on their screen and what the recruiter ultimately receives.",
            "When a user triggers the 'Download PDF' action, the system executes a complex sequence of events. First, the application state is temporarily locked to prevent any concurrent data mutations during the rendering process. The `react-to-print` engine then creates a hidden, isolated `iframe` within the browser's memory. It deeply clones the target `div` containing the resume template and injects it into this iframe. Crucially, the engine also traverses the primary document to extract all loaded stylesheets, Google Fonts, and dynamically injected styled-components, copying them into the iframe to ensure that all stylistic rules are preserved in the isolated printing environment.",
            "To guarantee absolute precision in the final output, the application utilizes highly specific `@media print` CSS rules. These rules explicitly define the page dimensions (exactly 210mm x 297mm for A4 sizing), remove all browser-injected margins, and disable background graphics printing defaults that browsers normally enforce. Furthermore, complex logic is implemented using the `break-inside: avoid` CSS property on critical components like experience blocks or education entries. This ensures that the browser's PDF rendering engine does not awkwardly slice a single job description across two separate pages, maintaining the professional integrity and readability of the final multi-page document."
        ]
    },
    {
        "title": "9. STATE MANAGEMENT AND DATA PERSISTENCE STRATEGY",
        "content": [
            "Given the highly interactive nature of the AI Resume Maker, efficient and robust state management is paramount. The application deliberately avoids the overhead and boilerplate of external state management libraries like Redux or MobX. Instead, it leverages React's native `useState` and `useEffect` hooks, combined with a custom-built context provider, to orchestrate the complex, deeply nested data structure of a professional resume. The root state object contains multiple arrays for experiences, education, and skills, alongside nested objects for personal information, resulting in a complex data tree that must be updated with sub-millisecond latency to ensure a smooth typing experience.",
            "To achieve this required performance, the application heavily utilizes functional state updates and immutable data patterns. When a user types into an input field deep within an 'Experience' component, a highly targeted callback function is triggered. This function does not mutate the original state object; rather, it uses the JavaScript spread operator (`...`) to create a shallow copy of the entire resume tree, updating only the specific string value that changed. This immutable approach guarantees that React's reconciliation engine can quickly and efficiently determine exactly which DOM nodes need to be re-rendered in the Preview pane, avoiding catastrophic, application-wide re-renders that would cause severe input lag.",
            "Data persistence is handled through a sophisticated dual-layer strategy. For immediate, session-level persistence, every state mutation triggers a `useEffect` hook that serializes the entire resume object into a JSON string and stores it synchronously in the browser's `sessionStorage`. This provides a crucial safety net, ensuring that if the user accidentally refreshes the page or closes their browser tab, their meticulously entered data is instantly recovered upon return. For long-term, cross-device persistence, the user can explicitly trigger a 'Save to Cloud' action. This asynchronous operation dispatches the JSON payload to the Next.js API route, which authenticates the user via their JWT token and executes a secure 'upsert' operation into the Supabase PostgreSQL database, ensuring their data is permanently and securely archived."
        ]
    },
    {
        "title": "10. SECURITY CONSIDERATIONS AND ETHICAL IMPLICATIONS",
        "content": [
            "In an application designed to collect, process, and store highly sensitive Personally Identifiable Information (PII)—including full names, contact details, comprehensive employment histories, and educational backgrounds—security cannot be an afterthought; it must be a foundational architectural principle. The AI Resume Maker addresses these critical security requirements at multiple layers of the stack. All network communication between the client browser, the Next.js serverless functions, and the Supabase database is strictly enforced over Transport Layer Security (TLS 1.3), ensuring that data is fully encrypted in transit and protected against Man-in-the-Middle (MitM) interception attacks.",
            "At the database layer, security is deeply integrated into the PostgreSQL engine using Row Level Security (RLS). Traditional applications rely on the application code to filter data and ensure users only see their own records. This approach is prone to catastrophic failure if a single API endpoint is misconfigured. In contrast, RLS pushes the security authorization logic directly into the database kernel. A custom policy is applied to the `resumes` table stating: `USING (auth.uid() = user_id)`. This guarantees at the cryptographic level that, regardless of the API query constructed by the client, the database will utterly refuse to return or mutate any record that does not strictly belong to the currently authenticated user's JWT.",
            "Beyond technical security, the project actively considers the ethical implications of automated resume generation. As Applicant Tracking Systems (ATS) become increasingly prevalent, filtering out candidates based on algorithmic keyword matching, there is a distinct risk of perpetuating systemic biases. The AI Resume Maker seeks to level the playing field. By providing democratized access to ATS-optimized, structurally flawless templates, the application ensures that candidates are judged purely on the merit of their experience and skills, rather than their ability to manipulate complex word processor formatting. Furthermore, strict data privacy policies are enforced, granting users complete autonomy and the absolute 'Right to Erasure', allowing them to permanently purge their digital footprint from the database at their discretion."
        ]
    }
]

def main():
    doc = Document()
    add_page_borders(doc)
    
    # Title Page
    doc.add_heading("AI RESUME MAKER", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("COMPREHENSIVE INTERNSHIP REPORT", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n\n\n\n")
    p = doc.add_paragraph("An in-depth technical analysis of the architecture, design, and implementation of a modern, scalable web application.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for idx, section in enumerate(content_sections):
        # Add heading
        add_heading(doc, section['title'], level=1)
        
        # Add paragraphs
        for para_text in section['content']:
            add_paragraph(doc, para_text)
            
        # Add image if exists
        if 'image' in section and os.path.exists(section['image']):
            try:
                # Add a little space
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run()
                r.add_picture(section['image'], width=Inches(6.0))
            except Exception as e:
                print(f"Error adding image {section['image']}: {e}")
                
        # Force a page break after each section to ensure "one topic, one page"
        # unless it's the very last section.
        if idx < len(content_sections) - 1:
            doc.add_page_break()

    output_path = r"C:\Users\Bhoomika\Desktop\AI_Resume_Maker_Final_Master_Report_V2.docx"
    doc.save(output_path)
    print(f"Massive DOCX with borders created successfully at: {output_path}")

if __name__ == "__main__":
    main()
