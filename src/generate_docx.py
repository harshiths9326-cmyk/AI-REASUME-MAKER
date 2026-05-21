from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_docx(input_text_file, docx_file):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(input_text_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    
    for line in lines:
        line_stripped = line.strip()
        
        if line_stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = doc.styles['No Spacing']
            p.paragraph_format.left_indent = Inches(0.5)
            # Use a monospaced font for code if possible
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            continue

        if not line_stripped:
            continue
            
        if line_stripped.startswith('# '):
            doc.add_page_break()
            p = doc.add_heading(line_stripped[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line_stripped.startswith('## '):
            doc.add_heading(line_stripped[3:], level=1)
        elif line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:], level=2)
        elif line_stripped.startswith('[IMAGE:'):
            parts = line_stripped[7:-1].split('|')
            img_path = parts[0].strip()
            caption = parts[1].strip() if len(parts) > 1 else ""
            
            brain_dir = r"C:\Users\Bhoomika\.gemini\antigravity\brain\7729b00b-8521-4b12-ad09-07549338132e"
            potential_path = os.path.join(brain_dir, img_path)
            if os.path.exists(potential_path):
                img_path = potential_path

            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6))
                    if caption:
                        cp = doc.add_paragraph(caption)
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in cp.runs:
                            run.font.italic = True
                            run.font.size = Pt(9)
                except Exception as e:
                    doc.add_paragraph(f"[Error loading image: {img_path}]")
            else:
                doc.add_paragraph(f"[Image not found: {img_path}]")
        else:
            if line_stripped.startswith('- '):
                p = doc.add_paragraph(line_stripped[2:], style='List Bullet')
            else:
                p = doc.add_paragraph(line_stripped)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(docx_file)
    print(f"DOCX created: {docx_file}")

if __name__ == "__main__":
    create_docx("ultimate_report_content.txt", "AI_Resume_Maker_Final_Master_Report.docx")
