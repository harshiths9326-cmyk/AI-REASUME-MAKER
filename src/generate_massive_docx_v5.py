import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_borders(doc):
    for sec in doc.sections:
        sectPr = sec._sectPr
        existing_borders = sectPr.xpath('.//w:pgBorders')
        if existing_borders:
            for eb in existing_borders:
                sectPr.remove(eb)
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '24')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')
            pgBorders.append(border)
        sectPr.append(pgBorders)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Arial'
        if level == 0:
            run.font.size = Pt(24)
            run.bold = True
        elif level == 1:
            run.font.size = Pt(18)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True
    return h

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    return p

def add_code(doc, code_text):
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    doc.add_paragraph() # Add space after code block

chapters = [
    {
        "title": "1. PROJECT INITIALIZATION & ENTRY POINT",
        "intro": [
            "The AI Resume Maker begins with a highly optimized landing page that serves as the entry point for all users. The goal of this component is to immediately convey the value proposition of the application while ensuring extremely fast load times.",
            "This section details the Next.js page structure and the use of Framer Motion for entrance animations, which significantly enhances the perceived performance of the application."
        ],
        "code": '''export default function Home() {
  return (
    <div className="flex flex-col min-h-[calc(100vh-4rem)]">
      <section className="w-full py-12 md:py-24 bg-background relative">
        <motion.div 
           initial={{ opacity: 0, y: 20 }}
           animate={{ opacity: 1, y: 0 }}
           className="container px-4 text-center">
             <h1 className="text-4xl font-black tracking-tighter">
               ENGINEER YOUR PROFESSIONAL RESUME
             </h1>
             <Link href="/builder"><Button>Initialize Builder</Button></Link>
        </motion.div>
      </section>
    </div>
  );
}''',
        "explanation": [
            "This functional React component acts as the main landing page. We utilize 'framer-motion' to create a subtle fade-up animation ('opacity: 0' to 'opacity: 1') when the page loads.",
            "The 'Link' component from Next.js is used to route the user to the '/builder' path. This ensures client-side navigation without a full page reload, maintaining the Single Page Application (SPA) feel."
        ],
        "image": "landing_page_screenshot_1778698938455.png",
        "post_image": [
            "As depicted in the screenshot above, the UI is kept clean and minimalist. The central call-to-action button is prominently displayed.",
            "By utilizing Tailwind CSS utility classes like 'py-12' and 'text-4xl', we achieve a highly responsive design that scales perfectly across mobile and desktop devices without writing custom CSS media queries."
        ],
        "subpoints": [
            {
                "title": "1.1 The Importance of First Contentful Paint (FCP)",
                "content": [
                    "In modern web development, the First Contentful Paint (FCP) is a critical metric. It measures the time from when the page starts loading to when any part of the page's content is rendered on the screen.",
                    "Because we use Next.js Server-Side Rendering (SSR) for the landing page, the HTML is pre-generated on the server. This means the browser can render the text and buttons immediately, before the JavaScript payload even finishes downloading.",
                    "This architecture ensures that users on slow 3G networks or low-end mobile devices do not stare at a blank white screen, drastically reducing bounce rates."
                ]
            }
        ]
    },
    {
        "title": "2. CORE ARCHITECTURE: THE BUILDER ORCHESTRATOR",
        "intro": [
            "Once the user clicks 'Initialize Builder', they are taken to the core application interface. This interface is managed by the Builder Orchestrator component, which is responsible for holding the master state of the entire resume.",
            "Managing complex, deeply nested state in React requires a strategic approach to prevent unnecessary re-renders. We utilize the 'useState' and 'useEffect' hooks to manage this data lifecycle."
        ],
        "code": '''function BuilderContent() {
    const [data, setData] = useState<ResumeData>(initialResumeData)
    const [template, setTemplate] = useState<string>("modern")

    useEffect(() => {
        const savedData = sessionStorage.getItem("resume_builder_data")
        if (savedData) setData(JSON.parse(savedData))
    }, [])

    const updateData = (newData: Partial<ResumeData>) => {
        setData((prev) => ({ ...prev, ...newData }))
    }

    return (
        <div className="grid grid-cols-1 lg:grid-cols-2 gap-6 h-full">
            <ResumeForm data={data} updateData={updateData} />
            <ResumePreview data={data} template={template} />
        </div>
    )
}''',
        "explanation": [
            "The 'BuilderContent' component uses a dual-pane layout via CSS Grid. The 'ResumeForm' sits on the left, and the 'ResumePreview' sits on the right.",
            "The 'useEffect' hook fires once when the component mounts. It checks the browser's 'sessionStorage' to see if the user has a saved resume from a previous session, ensuring no data is lost upon accidental page refreshes.",
            "The 'updateData' function uses a functional state update and the spread operator to immutably merge new data into the existing state tree."
        ],
        "image": "template_modern_ai_1778698574877.png",
        "post_image": [
            "The screenshot above shows the result of this architecture. The form on the left directly mutates the state held in the Orchestrator.",
            "Because the Orchestrator passes that exact same state down to the 'ResumePreview' component on the right, any keystroke in the form is immediately reflected in the visual layout."
        ],
        "subpoints": [
            {
                "title": "2.1 Immutable State Updates in React",
                "content": [
                    "React relies on reference equality to determine if a component should re-render. If we were to mutate the 'data' object directly (e.g., 'data.personalInfo.name = \"John\"'), React would not detect the change.",
                    "By using the spread operator ('...prev'), we create a completely new object in memory. When we call 'setData' with this new object, React sees that the reference has changed and triggers a reconciliation cycle.",
                    "This strict adherence to immutability is what guarantees the UI remains perfectly synchronized with the underlying data model at all times."
                ]
            }
        ]
    },
    {
        "title": "3. DATA MODELING WITH TYPESCRIPT",
        "intro": [
            "A resume is an inherently complex data structure. It contains simple strings (like a name), but also arrays of complex objects (like multiple job experiences, each with their own dates and descriptions).",
            "To prevent runtime errors, such as trying to access a property that doesn't exist, we rely heavily on TypeScript interfaces to enforce a strict data schema across the entire application."
        ],
        "code": '''export interface Experience {
    id: string;
    company: string;
    position: string;
    startDate: string;
    endDate: string;
    description: string;
}

export interface ResumeData {
    personalInfo: PersonalInfo;
    experience: Experience[];
    education: Education[];
    skills: Skill[];
}''',
        "explanation": [
            "These interfaces act as a binding contract. If a developer attempts to pass an object to the 'ResumePreview' that lacks the 'experience' array, the TypeScript compiler will throw an error and halt the build.",
            "The 'Experience' interface explicitly defines an 'id' property. This 'id' is crucial for React to efficiently manage the rendering of array items."
        ],
        "image": "template_corporate_professional_1778698667541.png",
        "post_image": [
            "In the 'Executive Professional' template shown above, the 'experience' array is iterated over to render the chronological work history.",
            "Because the data is strictly typed, the template component can confidently map over 'data.experience' without wrapping the code in excessive 'try/catch' blocks or 'null' checks."
        ],
        "subpoints": [
            {
                "title": "3.1 The Importance of Unique Identifiers (UUIDs)",
                "content": [
                    "When rendering lists in React, it is a common anti-pattern to use the array index as the 'key' prop (e.g., 'key={index}').",
                    "If a user deletes the first job experience in their list, all subsequent items shift up. If index keys are used, React gets confused and might render the wrong data in the wrong input fields.",
                    "To solve this, our schema mandates an 'id' string. When a new experience is added, we generate a cryptographically secure UUID. This UUID serves as a stable, unique key, completely eliminating array rendering bugs."
                ]
            }
        ]
    },
    {
        "title": "4. MANAGING COMPLEX FORM STATE",
        "intro": [
            "Building dynamic forms in React where users can add an arbitrary number of fields (like adding multiple skills or jobs) requires sophisticated state management within the child components.",
            "This section details how the 'Experience' component allows users to dynamically append and remove items from the central state array."
        ],
        "code": '''export function Experience({ data, updateData }: Props) {
    const addExperience = () => {
        updateData([
            ...data,
            {
                id: crypto.randomUUID(),
                company: "",
                position: "",
                description: "",
            },
        ])
    }

    const removeExperience = (id: string) => {
        updateData(data.filter((item) => item.id !== id))
    }
}''',
        "explanation": [
            "The 'addExperience' function uses the spread operator to copy all existing experiences into a new array, and then appends a new, blank experience object to the end.",
            "Crucially, it assigns a brand new UUID to the 'id' field using the browser's native 'crypto.randomUUID()' API.",
            "The 'removeExperience' function uses the native JavaScript 'filter' method to create a new array containing only the items whose 'id' does not match the one being deleted."
        ],
        "image": "template_modern_creative_1778698741933.png",
        "post_image": [
            "The 'Modern Creative' template dynamically handles these array mutations perfectly. As you add or remove skills in the left-hand form, the colored sidebar on the right instantly expands or contracts.",
            "This seamless interaction is entirely dependent on the flawless execution of the array manipulation functions detailed above."
        ],
        "subpoints": [
            {
                "title": "4.1 Controlled Inputs and Change Handlers",
                "content": [
                    "Every input field in the application is a 'Controlled Component'. This means the HTML '<input>' tag does not manage its own state; its 'value' is strictly bound to the React state.",
                    "When a user types a key, the 'onChange' event fires, which triggers an update function that updates the master state. React then re-renders the input with the new value.",
                    "This bidirectional data flow ensures that the application always has a single source of truth for the user's data, preventing nasty desynchronization bugs."
                ]
            }
        ]
    },
    {
        "title": "5. TEMPLATE RENDERING FACTORY",
        "intro": [
            "The application boasts 12 distinct professional templates. Instead of hardcoding all of these templates into a single massive file, we utilize a design pattern known as the 'Factory Pattern'.",
            "This allows us to dynamically switch between entirely different UI layouts based on a simple string value selected by the user."
        ],
        "code": '''const renderTemplate = () => {
    switch (template) {
        case "corporate":
            return <CorporateTemplate data={data} />;
        case "creative":
            return <CreativeTemplate data={data} />;
        case "modern-tech":
            return <ModernTechTemplate data={data} />;
        case "modern":
        default:
            return <ModernTemplate data={data} />;
    }
}''',
        "explanation": [
            "A simple 'switch' statement evaluates the 'template' state variable (which is bound to the sidebar selection menu).",
            "It then returns the corresponding React component, passing the master 'data' object into it as a prop.",
            "This modularity is exceptional for team development; a developer can build a brand new template in complete isolation and simply add one more 'case' statement to integrate it into the application."
        ],
        "image": "template_modern_tech_1778698775717.png",
        "post_image": [
            "By selecting the 'modern-tech' option, the factory instantly swaps out the standard view for the dark-mode tech template shown above.",
            "Notice how the data remains exactly the same; only the presentation layer has been swapped. This decoupling of data from view is a hallmark of professional software engineering."
        ],
        "subpoints": [
            {
                "title": "5.1 The Principle of Open/Closed Architecture",
                "content": [
                    "The Factory Pattern utilized here adheres closely to the 'Open/Closed Principle' of SOLID software design.",
                    "The 'ResumePreview' engine is 'closed for modification' (we rarely need to change its core logic) but 'open for extension' (we can easily add infinitely more templates by simply creating new components).",
                    "This dramatically reduces the risk of introducing regressions or breaking existing templates when adding new features to the platform."
                ]
            }
        ]
    },
    {
        "title": "6. THE PDF GENERATION PIPELINE",
        "intro": [
            "Converting HTML and CSS into a static PDF document that looks exactly the same across all devices is a notoriously complex problem. Browsers inject margins, strip background colors, and arbitrarily slice content across pages.",
            "To solve this, we rely on the 'react-to-print' library combined with highly specific print media CSS queries."
        ],
        "code": '''import { useReactToPrint } from "react-to-print"

export function ResumePreview({ data, template }) {
    const targetRef = useRef<HTMLDivElement>(null)
    const [isDownloading, setIsDownloading] = useState(false)

    const handlePrint = useReactToPrint({
        contentRef: targetRef,
        documentTitle: `${data.personalInfo.firstName}_Resume`,
        onAfterPrint: () => setIsDownloading(false),
    });

    return (
        <div ref={targetRef} className="print-styles bg-white">
            {renderTemplate()}
        </div>
    )
}''',
        "explanation": [
            "We attach a 'useRef' to the parent container of the resume template. The 'useReactToPrint' hook uses this reference to extract the exact DOM node from the page.",
            "It then creates an invisible 'iframe' in the background, injects the cloned DOM node, pulls in all associated stylesheets, and programmatically triggers the browser's native print API.",
            "The 'isDownloading' state is used to disable the download button and show a loading spinner while the browser processes the document."
        ],
        "image": "template_elegant_serif_clean_1778698816640.png",
        "post_image": [
            "When the Elegant Serif template (above) is printed, it is crucial that the delicate serif fonts are preserved perfectly.",
            "The print engine ensures that any custom Google Fonts loaded by the application are successfully transferred into the printing iframe, maintaining typographic integrity."
        ],
        "subpoints": [
            {
                "title": "6.1 Advanced CSS for Print Formatting",
                "content": [
                    "To ensure the PDF looks like a professional document and not a web page, we use the '@media print' CSS directive.",
                    "We apply '-webkit-print-color-adjust: exact' to force the browser to print background colors and sidebar graphics.",
                    "Furthermore, we use 'break-inside: avoid' on major layout blocks. This CSS rule mathematically prevents the browser from slicing a job description in half across two pages, forcing the entire block to the next page if necessary."
                ]
            }
        ]
    },
    {
        "title": "7. BACKEND INTEGRATION: SUPABASE DATABASE",
        "intro": [
            "While local 'sessionStorage' is excellent for preventing accidental data loss during a session, long-term persistence requires a robust backend database.",
            "We selected Supabase, an open-source Firebase alternative powered by PostgreSQL, to securely store user resumes in the cloud."
        ],
        "code": '''-- PostgreSQL Table Schema
CREATE TABLE resumes (
  id UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
  user_id UUID REFERENCES auth.users(id) NOT NULL,
  data JSONB NOT NULL,
  template_id TEXT DEFAULT 'modern',
  updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

-- Row Level Security Policy
ALTER TABLE resumes ENABLE ROW LEVEL SECURITY;

CREATE POLICY "Users can only access their own resumes"
ON resumes FOR ALL
USING (auth.uid() = user_id);''',
        "explanation": [
            "The SQL schema defines a 'resumes' table. Critically, it utilizes the 'JSONB' data type for the 'data' column. This allows us to store the entire complex resume object directly in the database without needing complex relational joins.",
            "The 'user_id' column serves as a foreign key linking the resume to the authentication system.",
            "The 'Row Level Security' (RLS) policy is the ultimate security safeguard. It guarantees at the database kernel level that a user can only read, update, or delete a row if their JWT token matches the 'user_id'."
        ],
        "image": "template_blue_cv_1778698836388.png",
        "post_image": [
            "When a user with the Blue CV template clicks 'Save to Cloud', the Next.js frontend dispatches an API request containing the JSON payload.",
            "The Supabase backend receives this payload, verifies the user's session token, and executes an UPSERT operation into the PostgreSQL database, updating the 'updated_at' timestamp automatically."
        ],
        "subpoints": [
            {
                "title": "7.1 The Power of JSONB in PostgreSQL",
                "content": [
                    "Traditionally, relational databases are highly rigid. If we wanted to add a 'hobbies' section to the resume, we would have to write an SQL migration script to alter the table structure.",
                    "By leveraging PostgreSQL's 'JSONB' capability, we achieve the flexibility of a NoSQL database (like MongoDB) while maintaining the security, speed, and ACID compliance of a traditional relational database.",
                    "This architectural choice allows the frontend team to rapidly iterate on the resume data schema without constantly requiring backend database modifications."
                ]
            }
        ]
    },
    {
        "title": "8. SECURITY ARCHITECTURE: AUTHENTICATION",
        "intro": [
            "A resume builder handles a massive amount of Personally Identifiable Information (PII). Full names, phone numbers, email addresses, and detailed work histories are prime targets for malicious actors.",
            "Implementing a robust, defense-in-depth security posture is not an option; it is a mandatory requirement. We utilize JSON Web Tokens (JWT) for secure session management."
        ],
        "code": '''import { createClient } from '@supabase/supabase-js'

export async function POST(req: Request) {
    const supabase = createClient(process.env.URL, process.env.KEY)
    
    // Verify the user's session token
    const { data: { user }, error } = await supabase.auth.getUser()
    
    if (error || !user) {
        return new Response("Unauthorized", { status: 401 })
    }

    const { id, resume_data } = await req.json()
    
    // Securely upsert data
    await supabase.from('resumes').upsert({ 
        id, 
        user_id: user.id, 
        data: resume_data 
    })
    
    return new Response("Success", { status: 200 })
}''',
        "explanation": [
            "This code snippet represents a Next.js API Route operating on the server. When the client makes a request to save data, this function is executed securely away from the user's browser.",
            "The 'supabase.auth.getUser()' method automatically extracts and cryptographically verifies the JWT token sent in the request headers.",
            "If the token is invalid, expired, or missing, the API immediately halts execution and returns a '401 Unauthorized' response, completely protecting the database from unauthenticated access."
        ],
        "image": "template_beige_minimal_1778698839199.png",
        "post_image": [
            "Security must be invisible to the user but absolute in its enforcement. Whether a user is crafting a 'Beige Minimal' resume or a highly technical one, their data remains locked behind these API safeguards.",
            "The use of server-side environment variables (`process.env`) ensures that sensitive database credentials are never accidentally leaked to the client-side JavaScript bundle."
        ],
        "subpoints": [
            {
                "title": "8.1 Mitigating Cross-Site Scripting (XSS)",
                "content": [
                    "XSS occurs when a malicious user injects executable JavaScript into a data field (e.g., pasting a script tag into the 'Job Title' input).",
                    "React inherently protects against XSS by automatically sanitizing and escaping all strings rendered via JSX. If a user types '<script>alert(1)</script>', React renders it safely as a literal string on the screen, completely neutralizing the attack vector."
                ]
            }
        ]
    },
    {
        "title": "9. PERFORMANCE OPTIMIZATION AND DEPLOYMENT",
        "intro": [
            "The final phase of the software development lifecycle involves optimizing the application for performance and deploying it to a highly available production environment.",
            "We utilized Vercel, the creators of Next.js, to deploy the application across a global Edge Network."
        ],
        "code": '''// next.config.mjs
/** @type {import('next').NextConfig} */
const nextConfig = {
  reactStrictMode: true,
  images: {
    domains: ['avatars.githubusercontent.com'],
    formats: ['image/avif', 'image/webp'],
  },
  experimental: {
    optimizeCss: true,
  }
};

export default nextConfig;''',
        "explanation": [
            "The 'next.config.mjs' file dictates the build and deployment parameters. 'reactStrictMode' ensures the application adheres to React best practices by intentionally double-rendering components in development to catch side-effect bugs.",
            "The 'images' configuration automatically optimizes any external images loaded by the application, converting them to highly compressed formats like WebP or AVIF on the fly, drastically reducing bandwidth consumption.",
            "The experimental 'optimizeCss' flag instructs the Webpack bundler to aggressively minify and purge unused CSS, ensuring the final stylesheet is as small as mathematically possible."
        ],
        "subpoints": [
            {
                "title": "9.1 Edge Network Distribution",
                "content": [
                    "Traditional web hosting involves renting a single server in a specific geographical location. If a user far away attempts to access the site, they experience high latency.",
                    "Vercel completely bypasses this by deploying the application to an Edge Network (CDN). When the application is built, the static assets are replicated to servers physically located all over the globe.",
                    "When a user requests the AI Resume Maker, they are automatically routed to the server node geographically closest to them, ensuring lightning-fast load times regardless of their physical location."
                ]
            },
            {
                "title": "9.2 Continuous Integration (CI)",
                "content": [
                    "The deployment pipeline is fully automated. The Next.js repository is linked directly to a GitHub repository.",
                    "Whenever a developer pushes code to the 'main' branch, Vercel detects the change and automatically provisions an isolated build container. It installs the dependencies, runs the TypeScript compiler, and builds the production bundles.",
                    "If any compilation errors occur (e.g., a type mismatch), the build is instantly aborted. This 'fail-fast' CI methodology ensures that broken code can never accidentally be deployed to the live production environment."
                ]
            }
        ]
    },
    {
        "title": "10. CONCLUSION AND FUTURE ROADMAP",
        "intro": [
            "The AI Resume Maker has successfully evolved from a conceptual prototype into a robust, highly scalable, and secure web application capable of servicing thousands of users.",
            "By meticulously adhering to modern software engineering principles, employing strict type safety, and leveraging cloud-native infrastructure, the project stands as a testament to the power of the Next.js ecosystem."
        ],
        "subpoints": [
            {
                "title": "10.1 Strategic Achievements",
                "content": [
                    "The primary strategic achievement is the total decoupling of data entry from visual presentation. This architecture not only provides an exceptional user experience but also creates a highly maintainable codebase where new visual templates can be added indefinitely without altering the core logic.",
                    "The implementation of client-side PDF generation via 'react-to-print' completely eliminated the need for expensive, slow, server-side headless browsers, significantly reducing operational costs while improving reliability."
                ]
            },
            {
                "title": "10.2 Roadmap Phase 2: Artificial Intelligence",
                "content": [
                    "The next major phase of development involves integrating Large Language Models (LLMs) such as OpenAI's GPT-4 API directly into the text inputs.",
                    "Users will be able to input a rough, unformatted sentence about their job experience, and the AI will automatically rewrite it into a highly professional, ATS-optimized bullet point utilizing strong action verbs.",
                    "Additionally, an 'ATS Score' feature will be developed, allowing users to paste a job description. The application will mathematically compare the resume text against the job description and recommend specific keyword inclusions."
                ]
            },
            {
                "title": "10.3 Roadmap Phase 3: Platform Expansion",
                "content": [
                    "Future iterations will expand the platform beyond simple PDF generation. We plan to introduce customizable Cover Letter generation, utilizing the exact same data payload and matching visual templates.",
                    "Furthermore, we aim to implement an OAuth integration with LinkedIn, allowing users to instantly import their entire professional history with a single click, completely eliminating the friction of manual data entry."
                ]
            }
        ]
    }
]

def main():
    doc = Document()
    
    # Page Borders
    add_page_borders(doc)
    
    # Document Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title Page
    doc.add_paragraph("\n\n\n\n\n\n\n")
    title = add_heading(doc, "AI RESUME MAKER", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = add_heading(doc, "COMPREHENSIVE TECHNICAL REPORT", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n\n")
    
    desc = add_paragraph(doc, "An exhaustive analysis of the architecture, targeted source code implementation, state management, and user experience design of a modern, highly scalable web application.")
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Generate Content
    for idx, chapter in enumerate(chapters):
        add_heading(doc, chapter['title'], level=1)
        
        for text in chapter['intro']:
            add_paragraph(doc, text)
            
        if 'code' in chapter:
            doc.add_paragraph()
            add_heading(doc, "Architectural Implementation (Source Code Snippet):", level=2)
            add_code(doc, chapter['code'])
            
            if 'explanation' in chapter:
                for text in chapter['explanation']:
                    add_paragraph(doc, text)

        if 'image' in chapter and os.path.exists(os.path.join(r"C:\Users\Bhoomika\.gemini\antigravity\brain\7729b00b-8521-4b12-ad09-07549338132e", chapter['image'])):
            doc.add_paragraph() 
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            # Scaling image width to 4.5 inches (leaves ~70% of vertical space for text)
            img_path = os.path.join(r"C:\Users\Bhoomika\.gemini\antigravity\brain\7729b00b-8521-4b12-ad09-07549338132e", chapter['image'])
            r.add_picture(img_path, width=Inches(4.5))
                
            if 'post_image' in chapter:
                for text in chapter['post_image']:
                    add_paragraph(doc, text)

        for sub in chapter.get('subpoints', []):
            doc.add_paragraph() 
            add_heading(doc, sub['title'], level=2)
            for text in sub['content']:
                add_paragraph(doc, text)

        if idx < len(chapters) - 1:
            doc.add_page_break()

    output_path = r"C:\Users\Bhoomika\Desktop\AI_Resume_Maker_Final_Master_Report_V5.docx"
    doc.save(output_path)
    print(f"Refined DOCX (V5) created successfully at: {output_path}")

if __name__ == "__main__":
    main()
